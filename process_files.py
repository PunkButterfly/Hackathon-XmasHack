import os
import PyPDF2
from PyPDF2 import PdfFileReader
import docx
import re
# import pandas as pd
import json
import aspose.words as aw
import traceback


def read_pdf_file(pdf_file_path: str) -> str:
    '''
        Открываем .pdf файл и, соединяя все страницы,
        возращем полный текст документа
        или None если не удалось открыть файл
    '''
    try:
        full_pdf_text = ''
        with open(pdf_file_path, "rb") as filehandle:
            pdf = PdfFileReader(filehandle)
            print(pdf)
            num_pages = pdf.getNumPages()
            for num in range(num_pages):
                page = pdf.getPage(num)
                full_pdf_text = full_pdf_text + '\n' + page.extractText()
        return full_pdf_text
    except:
        print(f'Ошибка: Не удалось открыть файл: {pdf_file_path}\n {traceback.format_exc()}')
        return None


def read_docx_file(docx_file_path: str) -> str:
    '''
        Открываем .docx файл и, соединяя все абзацы,
        возращем полный текст документа
        или None если не удалось открыть файл
    '''
    try:
        doc = docx.Document(docx_file_path)
        all_paras = doc.paragraphs
        full_doc_text = ''
        for para in all_paras:
            full_doc_text = full_doc_text + '\n' + para.text
        return full_doc_text.replace('Evaluation Only. Created with Aspose.Words. Copyright 2003-2022 Aspose Pty Ltd.',
                                     '')
    except:
        print(f'Ошибка: Не удалось открыть файл: {docx_file_path}\n {traceback.format_exc()}')
        return None


def convert_doc_docx(doc_file_paths: str, converted_files_dir_path: str) -> list[tuple[str, str]]:
    '''
        Конвертируем .doc файлы в .docx, так как python не открывает файлы .doc
        Возвращаемые кортежи содержат правилльный полный путь к файлу, но старое расширение в имени файла
        Новые файлы содержат в начале Evaluation Only. Created with Aspose.Words. Copyright 2003-2022 Aspose Pty Ltd.
    '''
    new_docx_paths = []
    file_extension = re.compile(r'\..*')
    for doc_file_path, doc_file_name in doc_file_paths:
        doc = aw.Document(doc_file_path)
        new_docx_path = os.path.join(converted_files_dir_path, file_extension.sub("", doc_file_name) + '.docx')
        doc.save(new_docx_path)
        new_docx_paths.append((new_docx_path, doc_file_name))
    return new_docx_paths


def convert_file_to_text(document_file):
    doc_type = re.findall(r"\..*", document_file.name)[0]
    if doc_type == ".docx":
        document_content = read_docx_file(document_file)
    elif doc_type == ".pdf":
        with open("files/file.pdf", 'wb') as f:
            f.write(document_file.getvalue())
        document_content = read_pdf_file("files/file.pdf")
    elif doc_type == '.doc' or doc_type == '.rtf':
        with open("files/file.doc", 'wb') as f:
            f.write(document_file.getvalue())
        doc = aw.Document("files/file.doc")
        doc.save("files/file.docx")
        document_content = read_docx_file("files/file.docx")

    return document_content
